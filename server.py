from flask import Flask, request, jsonify, send_from_directory, send_file
import msgpack
import os
import sys
from functools import lru_cache
import re
from io import BytesIO
from werkzeug.utils import secure_filename

try:
    from docx import Document
    from bs4 import BeautifulSoup
except ImportError:
    pass

# Ensure UTF-8
if sys.stdout.encoding != 'utf-8':
    try:
        sys.stdout.reconfigure(encoding='utf-8')
    except AttributeError:
        import codecs
        sys.stdout = codecs.getwriter("utf-8")(sys.stdout.detach())

app = Flask(__name__, static_folder='frontend')

def normalize_kikuyu(text):
    if not isinstance(text, str): return ""
    text = text.lower()
    text = text.replace('ĩ', 'i').replace('ũ', 'u')
    return text

class TrieNode:
    def __init__(self):
        self.children = {}
        self.actual_words = {}  
        self.max_freq_in_subtree = 0  

    @staticmethod
    def from_dict(d):
        node = TrieNode()
        node.actual_words = d['w']
        node.max_freq_in_subtree = d['m']
        for char, child_dict in d['c'].items():
            node.children[char] = TrieNode.from_dict(child_dict)
        return node

class KikuyuTrie:
    def __init__(self):
        self.root = TrieNode()

    @staticmethod
    def load(file_path):
        with open(file_path, 'rb') as f:
            data = msgpack.unpack(f, raw=False)
        trie = KikuyuTrie()
        trie.root = TrieNode.from_dict(data)
        return trie

    def insert(self, word):
        norm_word = normalize_kikuyu(word)
        node = self.root
        node.max_freq_in_subtree = max(node.max_freq_in_subtree, 1)
        for char in norm_word:
            if char not in node.children:
                node.children[char] = TrieNode()
            node = node.children[char]
            node.max_freq_in_subtree = max(node.max_freq_in_subtree, 1)
        if word not in node.actual_words:
            node.actual_words[word] = 1

    def contains(self, word):
        norm_word = normalize_kikuyu(word)
        node = self.root
        for char in norm_word:
            if char not in node.children: return False
            node = node.children[char]
        return len(node.actual_words) > 0

    def autocorrect(self, word):
        norm_word = normalize_kikuyu(word)
        node = self.root
        for char in norm_word:
            if char not in node.children: return None
            node = node.children[char]
        if not node.actual_words: return None
        if word in node.actual_words: return None
        best_word = max(node.actual_words.items(), key=lambda x: x[1])[0]
        if best_word == word: return None
        return best_word

    @lru_cache(maxsize=2048)
    def search(self, prefix, top_n=5):
        import heapq
        norm_prefix = normalize_kikuyu(prefix)
        node = self.root
        for char in norm_prefix:
            if char not in node.children: return []
            node = node.children[char]
        
        pq = [(-node.max_freq_in_subtree, 0, node)]
        tie_breaker = 0
        
        results_heap = []
        result_tie = 0
        
        while pq:
            neg_max_freq, _, curr_node = heapq.heappop(pq)
            max_freq = -neg_max_freq
            
            if len(results_heap) == top_n and max_freq <= results_heap[0][0]:
                continue
                
            if curr_node.actual_words:
                for word, freq in curr_node.actual_words.items():
                    if len(results_heap) < top_n:
                        heapq.heappush(results_heap, (freq, result_tie, word))
                        result_tie += 1
                    elif freq > results_heap[0][0]:
                        heapq.heappushpop(results_heap, (freq, result_tie, word))
                        result_tie += 1
                        
            for char, child_node in curr_node.children.items():
                if len(results_heap) < top_n or child_node.max_freq_in_subtree > results_heap[0][0]:
                    tie_breaker += 1
                    heapq.heappush(pq, (-child_node.max_freq_in_subtree, tie_breaker, child_node))
                    
        final_results = [(freq, word) for freq, tie, word in results_heap]
        final_results.sort(key=lambda x: x[0], reverse=True)
        return [word for freq, word in final_results]

# Load the models locally from the models/ folder
BASE_DIR = os.path.dirname(__file__)
TRIE_PATH = os.path.join(BASE_DIR, 'models', 'kikuyu_trie.msgpack')
BIGRAM_PATH = os.path.join(BASE_DIR, 'models', 'kikuyu_bigrams.msgpack')
USER_DICT_PATH = os.path.join(BASE_DIR, 'user_dictionary.txt')

print("Loading Kikuyu Autocomplete Models...")
trie = KikuyuTrie.load(TRIE_PATH)

if os.path.exists(USER_DICT_PATH):
    print("Loading user dictionary...")
    with open(USER_DICT_PATH, 'r', encoding='utf-8') as f:
        for line in f:
            w = line.strip()
            if w: trie.insert(w)

with open(BIGRAM_PATH, 'rb') as f:
    bigrams = msgpack.unpack(f, raw=False)
print("Standalone Web Server Ready.")

@app.route('/suggest', methods=['GET'])
def suggest():
    query = request.args.get('q', '').lower()
    if not query: return jsonify([])
    return jsonify(trie.search(query, top_n=5))

@app.route('/check_word', methods=['GET'])
def check_word():
    word = request.args.get('w', '').strip()
    if not word: return jsonify({"valid": True})
    if re.match(r'^[\d\W_]+$', word):
        return jsonify({"valid": True})
    return jsonify({"valid": trie.contains(word)})

@app.route('/learn_word', methods=['POST'])
def learn_word():
    data = request.get_json()
    if not data or 'word' not in data:
        return jsonify({"success": False}), 400
    
    word = data['word'].strip()
    if word and len(word) <= 40 and re.match(r"^[a-zA-ZĩũĨŨ\']+$", word):
        trie.insert(word)
        trie.search.cache_clear()
        with open(USER_DICT_PATH, 'a', encoding='utf-8') as f:
            f.write(word + '\n')
        return jsonify({"success": True})
    return jsonify({"success": False, "error": "Invalid word"}), 400

@lru_cache(maxsize=2048)
def get_cached_prediction(prev_word):
    if prev_word in bigrams:
        return bigrams[prev_word]
    norm_prev = normalize_kikuyu(prev_word)
    if norm_prev in bigrams:
        return bigrams[norm_prev]
    return []

@app.route('/predict', methods=['GET'])
def predict():
    prev_word = request.args.get('prev', '').lower()
    return jsonify(get_cached_prediction(prev_word))

@app.route('/autocorrect', methods=['GET'])
def autocorrect_endpoint():
    word = request.args.get('w', '').strip()
    if not word: return jsonify({"correction": None})
    if re.match(r'^[\d\W_]+$', word): return jsonify({"correction": None})
    correction = trie.autocorrect(word)
    return jsonify({"correction": correction})

# --- Document Management API ---
DOCS_DIR = os.path.join(BASE_DIR, 'documents')
os.makedirs(DOCS_DIR, exist_ok=True)

@app.route('/api/documents', methods=['GET'])
def list_documents():
    files = [f for f in os.listdir(DOCS_DIR) if f.endswith('.html') or f.endswith('.kikdoc')]
    return jsonify(files)

@app.route('/api/documents/<filename>', methods=['GET'])
def get_document(filename):
    if not (filename.endswith('.html') or filename.endswith('.kikdoc')): 
        return jsonify({"error": "Invalid file"}), 400
    safe_name = secure_filename(filename)
    if not safe_name or safe_name != filename:
        return jsonify({"error": "Invalid file"}), 400
    try:
        with open(os.path.join(DOCS_DIR, safe_name), 'r', encoding='utf-8') as f:
            return jsonify({"content": f.read()})
    except FileNotFoundError:
        return jsonify({"error": "Not found"}), 404

@app.route('/api/documents/<filename>', methods=['POST'])
def save_document(filename):
    if not (filename.endswith('.html') or filename.endswith('.kikdoc')): 
        return jsonify({"error": "Invalid file"}), 400
    safe_name = secure_filename(filename)
    if not safe_name or safe_name != filename:
        return jsonify({"error": "Invalid file"}), 400
    data = request.get_json()
    if not data or 'content' not in data:
        return jsonify({"error": "No content"}), 400
    with open(os.path.join(DOCS_DIR, safe_name), 'w', encoding='utf-8') as f:
        f.write(data['content'])
    return jsonify({"success": True})

@app.route('/api/documents/<filename>', methods=['DELETE'])
def delete_document(filename):
    if not (filename.endswith('.html') or filename.endswith('.kikdoc')): 
        return jsonify({"error": "Invalid file"}), 400

    safe_name = secure_filename(filename)
    if not safe_name or safe_name != filename:
        return jsonify({"error": "Invalid file"}), 400

    base_dir = os.path.abspath(DOCS_DIR)
    file_path = os.path.abspath(os.path.join(base_dir, safe_name))
    if os.path.commonpath([base_dir, file_path]) != base_dir:
        return jsonify({"error": "Invalid file path"}), 400

    try:
        if os.path.exists(file_path):
            os.remove(file_path)
            return jsonify({"success": True})
        else:
            return jsonify({"error": "Not found"}), 404
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/export/docx', methods=['POST'])
def export_docx():
    data = request.get_json()
    if not data or 'html' not in data:
        return jsonify({"error": "No HTML content"}), 400
        
    try:
        html_content = data['html']
        soup = BeautifulSoup(html_content, 'html.parser')
        doc = Document()
        
        # Super basic HTML to docx parser
        for element in soup.find_all(['h1', 'h2', 'p']):
            if element.name == 'h1':
                doc.add_heading(element.get_text(), 1)
            elif element.name == 'h2':
                doc.add_heading(element.get_text(), 2)
            elif element.name == 'p':
                p = doc.add_paragraph()
                for child in element.children:
                    if child.name is None: # text node
                        if child.string: p.add_run(child.string)
                    elif child.name in ['b', 'strong']:
                        p.add_run(child.get_text()).bold = True
                    elif child.name in ['i', 'em']:
                        p.add_run(child.get_text()).italic = True
                    elif child.name == 'u':
                        p.add_run(child.get_text()).underline = True
                    elif child.name == 'br':
                        p.add_run('\n')
                    else:
                        p.add_run(child.get_text())

        f = BytesIO()
        doc.save(f)
        f.seek(0)
        
        return send_file(
            f,
            as_attachment=True,
            download_name='Kikuyu_Document.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except Exception as e:
        print(f"Export error: {e}")
        return jsonify({"error": "Export failed"}), 500

@app.route('/')
def index():
    return send_from_directory('frontend', 'index.html')

@app.route('/<path:path>')
def static_proxy(path):
    return send_from_directory('frontend', path)

if __name__ == '__main__':
    debug_mode = os.environ.get('FLASK_ENV') == 'development'
    app.run(host='127.0.0.1', port=5001, debug=debug_mode)
